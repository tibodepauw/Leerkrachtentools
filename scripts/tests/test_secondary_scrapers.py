from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(SCRIPTS_DIR))

from secondary_minimum_goals_common import extract_api_goals, normalize_api_goal
from secondary_record_schema import (
    GCS_SEPARATOR,
    export_secundair_gcs,
    format_gcs_paragraph,
    normalize_curriculum_record,
    normalize_finaliteit,
    normalize_graad,
    normalize_minimum_goal_record,
    validate_secondary_record,
)
from fetch_pov_secondary_curricula import normalize_pov_goal
from scrape_secondary_curricula import (
    SourceDocument,
    infer_discipline,
    infer_grade,
    infer_stream,
    parse_kov_docx,
)


class SecondarySchemaTests(unittest.TestCase):
    def test_normalizes_graad_and_finaliteit(self) -> None:
        self.assertEqual(normalize_graad("eerste graad"), "1ste graad")
        self.assertEqual(normalize_graad("7de leerjaar"), "7de specialisatiejaar")
        self.assertEqual(normalize_finaliteit("doorstroomfinaliteit"), "Doorstroom")
        self.assertEqual(normalize_finaliteit("", "A-stroom"), "A-stroom")
        self.assertEqual(normalize_finaliteit("dubbele finaliteit"), "Dubbele finaliteit")

    def test_curriculum_record_has_flat_schema(self) -> None:
        record = normalize_curriculum_record(
            {
                "code": "BV.1.A.1",
                "titel": "De leerlingen lossen problemen op.",
                "discipline": "Wiskunde",
                "graad": "1ste graad",
                "finaliteit": "A-stroom",
                "stroom": "A-stroom",
                "netwerk": "GO",
                "bron_url": "https://example.test/plan.pdf",
            }
        )
        self.assertEqual(record["onderwijsniveau"], "SECUNDAIR")
        self.assertEqual(record["finaliteit"], "A-stroom")
        self.assertEqual(record["netwerk"], "GO")
        self.assertEqual(validate_secondary_record(record), [])

    def test_minimum_goal_record_is_flat(self) -> None:
        record = normalize_minimum_goal_record(
            code="06.12",
            text="De leerlingen analyseren een wiskundig probleem.",
            goal_type="Eindtermen",
            grade="1ste graad",
            stream="A-stroom",
            sc_name="Wiskunde",
            source_url="https://www.onderwijsdoelen.be/doelen/SO_1STE_GRAAD_V2_1",
        )
        self.assertEqual(record["code"], "06.12")
        self.assertEqual(
            record["titel"],
            "De leerlingen analyseren een wiskundig probleem.",
        )
        self.assertEqual(record["onderwijsniveau"], "SECUNDAIR")
        self.assertNotIn("gelinkt_minimumdoel", record)

    def test_gcs_paragraph_has_no_json_braces(self) -> None:
        paragraph = format_gcs_paragraph(
            normalize_minimum_goal_record(
                code="06.12",
                text="De leerlingen analyseren een wiskundig probleem.",
                sc_name="Wiskunde",
            )
        )
        self.assertNotIn("{", paragraph)
        self.assertNotIn("}", paragraph)
        self.assertIn("Doelcode: 06.12", paragraph)

    def test_exports_network_specific_gcs_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            data_dir = Path(tmp) / "secundair"
            data_dir.mkdir()
            curriculum = [
                normalize_curriculum_record(
                    {
                        "code": "K1",
                        "titel": "KOV doel",
                        "discipline": "Nederlands",
                        "graad": "2de graad",
                        "finaliteit": "Doorstroom",
                        "netwerk": "KOV",
                        "bron_url": "https://kov.test",
                    }
                ),
                normalize_curriculum_record(
                    {
                        "code": "G1",
                        "titel": "GO doel",
                        "discipline": "Wiskunde",
                        "graad": "1ste graad",
                        "finaliteit": "A-stroom",
                        "stroom": "A-stroom",
                        "netwerk": "GO",
                        "bron_url": "https://go.test",
                    }
                ),
            ]
            minimum = [
                normalize_minimum_goal_record(
                    code="06.12",
                    text="Minimumdoel",
                    sc_name="Wiskunde",
                )
            ]
            (data_dir / "leerplannen_secundair.jsonl").write_text(
                "\n".join(json.dumps(item, ensure_ascii=False) for item in curriculum)
                + "\n",
                encoding="utf-8",
            )
            (data_dir / "minimumdoelen_secundair.jsonl").write_text(
                json.dumps(minimum[0], ensure_ascii=False) + "\n",
                encoding="utf-8",
            )

            output_dir = data_dir / "gcs"
            counts = export_secundair_gcs(data_dir, output_dir)
            self.assertEqual(counts["leerplannen_secundair_kov.txt"], 1)
            self.assertEqual(counts["leerplannen_secundair_go.txt"], 1)
            self.assertEqual(counts["minimumdoelen_secundair.txt"], 1)

            kov_text = (output_dir / "leerplannen_secundair_kov.txt").read_text(
                encoding="utf-8"
            )
            self.assertIn(GCS_SEPARATOR, kov_text)
            self.assertIn("KOV doel", kov_text)


class SecondaryCurriculumParserTests(unittest.TestCase):
    def test_infers_secondary_routes(self) -> None:
        self.assertEqual(infer_grade("Onderwijsdoelen eerste graad"), "1ste graad")
        self.assertEqual(infer_grade("Leerplan 3de graad"), "3de graad")
        self.assertEqual(infer_stream("Artistieke vorming A+B-stroom"), "A- en B-stroom")
        self.assertEqual(
            infer_discipline(
                "pdf Leerplan Leer Lokaal SO - eerste graad - A-stroom - "
                "basisvorming - SC 06 Wiskunde 09/07/2025"
            ),
            "Wiskunde",
        )

    def test_parses_kov_goal_styles_and_minimum_links(self) -> None:
        from docx import Document
        from io import BytesIO

        document = Document()
        document.add_heading("Leerplandoelen", level=1)
        minimum = document.add_paragraph(
            "MD 09.02 De leerlingen beschrijven landschappen. (LPD 1)"
        )
        minimum.style = document.styles["Normal"]
        document.styles.add_style("MD + SMD + BK", 1)
        minimum.style = document.styles["MD + SMD + BK"]
        document.styles.add_style("Doel", 1)
        goal = document.add_paragraph(
            "De leerlingen beschrijven kenmerken van landschappen."
        )
        goal.style = document.styles["Doel"]

        payload = BytesIO()
        document.save(payload)
        source = SourceDocument(
            provider="KOV",
            title="Aardrijkskunde - A-stroom",
            url="https://example.test/aardrijkskunde.docx",
            page_url="https://pro.katholiekonderwijs.vlaanderen/I-Aar-a/leerplan",
            grade="1ste graad",
            stream="A-stroom",
            discipline="Aardrijkskunde",
            file_type="docx",
        )

        records = parse_kov_docx(payload.getvalue(), source)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0].code, "I-Aar-a LPD 1")
        self.assertEqual(records[0].minimumdoel_codes, ["09.02"])

        normalized = normalize_curriculum_record(records[0])
        self.assertEqual(normalized["finaliteit"], "A-stroom")
        self.assertEqual(normalized["onderwijsniveau"], "SECUNDAIR")


class SecondaryMinimumGoalsParserTests(unittest.TestCase):
    def test_normalizes_official_api_goal(self) -> None:
        result = normalize_api_goal(
            {
                "uniqueCode": "06.12",
                "title": "De leerlingen analyseren een wiskundig probleem.",
                "onderwijsniveau": "secundair onderwijs",
                "type": "eindterm",
                "graad": "eerste graad",
                "stroom": "A-stroom",
                "sleutelcompetentie": "Wiskunde",
            }
        )
        self.assertIsNotNone(result)
        assert result is not None
        self.assertEqual(result["code"], "06.12")
        self.assertEqual(result["graad"], "1ste graad")
        self.assertEqual(result["finaliteit"], "A-stroom")
        self.assertEqual(result["onderwijsniveau"], "SECUNDAIR")

    def test_normalizes_portal_goal_with_sc_metadata(self) -> None:
        from secondary_minimum_goals_common import normalize_portal_goal

        result = normalize_portal_goal(
            {
                "code": "06.12",
                "omschrijving": "<p>De leerlingen analyseren een wiskundig probleem.</p>",
                "onderwijsdoel_type": "Eindtermen",
                "_dataset": "SO_1STE_GRAAD_V2_1",
                "onderwijsdoelenset": {
                    "onderwijsdoelenset": "Secundair onderwijs 1ste graad A-stroom - Wiskunde - Eindtermen",
                    "vlaamse_sleutelcompetentie": {
                        "nr": "6",
                        "naam": "Wiskunde - natuurwetenschappen - technologie en techniek (STEM)",
                    },
                    "onderwijsstructuur": {
                        "graad": "1ste graad",
                        "stroom": "A-stroom",
                    },
                },
            }
        )
        self.assertIsNotNone(result)
        assert result is not None
        self.assertEqual(result["graad"], "1ste graad")
        self.assertEqual(result["finaliteit"], "A-stroom")
        self.assertEqual(result["discipline"], "Wiskunde - natuurwetenschappen - technologie en techniek (STEM)")

    def test_deduplicates_nested_api_payloads(self) -> None:
        goal = {
            "code": "02.01",
            "titel": "De leerlingen verwerken doelgericht informatie.",
            "onderwijsniveau": "secundair onderwijs",
        }
        results = extract_api_goals([{"items": [goal, {"nested": goal}]}])
        self.assertEqual(len(results), 1)


class PovCurriculumParserTests(unittest.TestCase):
    def test_normalizes_detailed_api_goal(self) -> None:
        result = normalize_pov_goal(
            {
                "code": "LPD 3",
                "doelzin": "De leerlingen onderzoeken eigenschappen van materialen.",
                "vak": "Techniek",
            },
            {
                "naam": "Techniek - 2de graad - dubbele finaliteit",
            },
        )
        self.assertIsNotNone(result)
        assert result is not None
        self.assertEqual(result["netwerk"], "POV")
        self.assertEqual(result["graad"], "2de graad")
        self.assertEqual(result["finaliteit"], "Dubbele finaliteit")
        self.assertEqual(result["onderwijsniveau"], "SECUNDAIR")


if __name__ == "__main__":
    unittest.main()
