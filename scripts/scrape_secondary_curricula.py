#!/usr/bin/env python3
"""
Download en normaliseer publieke Vlaamse leerplannen secundair onderwijs.

Bronnen:
  - GO! Pro: publieke leerplan-PDF's voor 1ste, 2de en 3de graad
  - Katholiek Onderwijs Vlaanderen: publieke Word-leerplannen
  - OVSG: publieke Leer Lokaal SO-documenten (eigen plannen vooral 1ste graad)

Uitvoer:
  data/secundair/leerplannen_secundair.jsonl

Gebruik:
  python3 scripts/scrape_secondary_curricula.py
  python3 scripts/scrape_secondary_curricula.py --provider GO --limit 2
  python3 scripts/scrape_secondary_curricula.py --skip-download
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import logging
import re
import sys
from collections import Counter
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "data/secundair/leerplannen_secundair.jsonl"
REPORT_PATH = OUTPUT_PATH.with_name("leerplannen_secundair_report.json")
DOWNLOAD_DIR = OUTPUT_PATH.parent / "downloads"

GO_BASE = "https://pro.g-o.be"
GO_PAGES = (
    ("1ste graad", "A-stroom", f"{GO_BASE}/themas/leerplannen/eerste-graad-secundair-onderwijs/a-stroom/"),
    ("1ste graad", "B-stroom", f"{GO_BASE}/themas/leerplannen/eerste-graad-secundair-onderwijs/b-stroom/"),
    ("2de graad", "doorstroomfinaliteit", f"{GO_BASE}/themas/leerplannen/tweede-graad-secundair-onderwijs/doorstroomfinaliteit/"),
    ("2de graad", "dubbele finaliteit", f"{GO_BASE}/themas/leerplannen/tweede-graad-secundair-onderwijs/dubbele-finaliteit/"),
    ("2de graad", "arbeidsmarktfinaliteit", f"{GO_BASE}/themas/leerplannen/tweede-graad-secundair-onderwijs/arbeidsmarktfinaliteit/"),
    ("3de graad", "doorstroomfinaliteit", f"{GO_BASE}/themas/leerplannen/derde-graad-secundair-onderwijs/doorstroomfinaliteit/"),
    ("3de graad", "dubbele finaliteit", f"{GO_BASE}/themas/leerplannen/derde-graad-secundair-onderwijs/dubbele-finaliteit/"),
    ("3de graad", "arbeidsmarktfinaliteit", f"{GO_BASE}/themas/leerplannen/derde-graad-secundair-onderwijs/arbeidsmarktfinaliteit/"),
)

KOV_BASE = "https://pro.katholiekonderwijs.vlaanderen"
KOV_INDEX = f"{KOV_BASE}/vakken-en-leerplannen"
KOV_TABS = (
    ("1ste graad", "eerstegraad"),
    ("2de graad", "tweedegraad"),
    ("3de graad", "derdegraad"),
    ("7de leerjaar", "zevendejaar"),
)

OVSG_BASE = "https://www.ovsg.be"
OVSG_INDEX = f"{OVSG_BASE}/leerplannen-secundair-onderwijs"

USER_AGENT = (
    "Leerkrachtentools-secondary-curriculum/1.0 "
    "(publieke onderwijsdata; https://github.com/tibodepauw/Leerkrachtentools)"
)
GOAL_CODE_RE = re.compile(
    r"(?m)^\s*((?:[AB]-[A-Z]{2,}\.\d+\.[A-Z]\.\d+)|"
    r"(?:(?:BV|BG|UD|VD|LPD|GO|SC)[A-Z0-9_-]*(?:\.\d+)+(?:\.\d+)*))\s*$"
)
MINIMUM_CODE_RE = re.compile(r"(?m)^\s*(?:MD\s*)?(\d{2}\.\d{2}(?:\.\d+)?)\s*$")
LEARNER_SENTENCE_RE = re.compile(
    r"(De (?:leerling|leerlingen)en?\b[\s\S]{15,900}?[.!?])"
)

logger = logging.getLogger("scrape_secondary_curricula")


def clean_text(value: Any) -> str:
    text = html.unescape(str(value or ""))
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def safe_name(value: str, suffix: str) -> str:
    name = re.sub(r"[^\w.-]+", "_", clean_text(value), flags=re.UNICODE).strip("_")
    digest = hashlib.sha1(value.encode("utf-8")).hexdigest()[:8]
    return f"{name[:100]}_{digest}{suffix}"


def infer_discipline(title: str) -> str:
    sc_match = re.search(
        r"\bSC\s*\d+\s+(.+?)(?:\s+\d{2}/\d{2}/\d{4})?$",
        clean_text(title),
        re.I,
    )
    if sc_match:
        return clean_text(sc_match.group(1))
    value = re.sub(
        r"(?i)^(?:pdf\s+)?(?:GO!\s+SO\s+)?(?:leerplan\s+)?"
        r"(?:Leer Lokaal SO\s*-\s*)?",
        "",
        clean_text(title),
    )
    value = re.sub(
        r"(?i)\b(?:1ste|2de|3de)\s+graad\b|"
        r"\b(?:A\+B|A|B)-stroom\b|"
        r"\b(?:doorstroom|dubbele|arbeidsmarkt)(?:finaliteit)?\b|"
        r"\bleerplan\b",
        " ",
        value,
    )
    return clean_text(value.strip(" -–"))


def infer_stream(text: str) -> str:
    lowered = text.casefold()
    if "a+b-stroom" in lowered:
        return "A- en B-stroom"
    if "a-stroom" in lowered:
        return "A-stroom"
    if "b-stroom" in lowered:
        return "B-stroom"
    return ""


def infer_grade(text: str) -> str:
    lowered = text.casefold()
    for pattern, label in (
        (r"\b(?:1ste|eerste)\s+graad\b", "1ste graad"),
        (r"\b(?:2de|tweede)\s+graad\b", "2de graad"),
        (r"\b(?:3de|derde)\s+graad\b", "3de graad"),
        (r"\b(?:7de|zevende)\s+(?:leerjaar|jaar)\b", "7de leerjaar"),
    ):
        if re.search(pattern, lowered):
            return label
    return ""


@dataclass(frozen=True)
class SourceDocument:
    provider: str
    title: str
    url: str
    page_url: str
    grade: str
    finality: str = ""
    stream: str = ""
    discipline: str = ""
    file_type: str = ""


@dataclass
class GoalRecord:
    code: str
    titel: str
    discipline: str
    subdomein: str
    toelichting: str
    leerjaar_route: str
    onderwijsniveau: str
    graad: str
    finaliteit: str
    stroom: str
    netwerk: str
    bron_url: str
    bron_titel: str
    minimumdoel_codes: list[str] = field(default_factory=list)
    sleutelcompetentie_nr: str = ""
    sleutelcompetentie: str = ""

    def key(self) -> tuple[str, str, str]:
        return (self.netwerk, self.code, self.titel.casefold())


class SecondaryCurriculumScraper:
    def __init__(
        self,
        *,
        output: Path = OUTPUT_PATH,
        providers: Iterable[str] = ("GO", "KOV", "OVSG"),
        limit: int | None = None,
        skip_download: bool = False,
        timeout: int = 90,
    ) -> None:
        self.output = output
        self.report_path = output.with_name("leerplannen_secundair_report.json")
        self.download_dir = output.parent / "downloads"
        self.providers = tuple(providers)
        self.limit = limit
        self.skip_download = skip_download
        self.timeout = timeout
        self.session = requests.Session()
        self.session.headers.update(
            {"User-Agent": USER_AGENT, "Accept": "text/html,application/pdf,*/*"}
        )
        self.warnings: list[str] = []

    def run(self) -> list[GoalRecord]:
        documents: list[SourceDocument] = []
        if "GO" in self.providers:
            documents.extend(self.discover_go())
        if "KOV" in self.providers:
            documents.extend(self.discover_kov())
        if "OVSG" in self.providers:
            documents.extend(self.discover_ovsg())

        documents = self._dedupe_documents(documents)
        if self.limit is not None:
            documents = documents[: self.limit]
        if not documents:
            raise RuntimeError("Geen publieke secundaire leerplandocumenten gevonden.")

        records: list[GoalRecord] = []
        parsed_documents = 0
        for index, source in enumerate(documents, start=1):
            logger.info("[%s/%s] %s", index, len(documents), source.title)
            try:
                payload, file_type = self._load_document(source)
                parsed = (
                    parse_kov_docx(payload, source)
                    if file_type == "docx"
                    else parse_pdf(payload, source)
                )
                if not parsed:
                    self._warn(f"Geen doelen herkend in {source.title}")
                    continue
                parsed_documents += 1
                records.extend(parsed)
            except Exception as exc:
                self._warn(f"{source.title}: {exc}")

        records = self._dedupe_records(records)
        if not records:
            raise RuntimeError("Documenten gevonden, maar geen leerplandoelen geëxtraheerd.")
        self._write(records, documents, parsed_documents)
        return records

    def discover_go(self) -> list[SourceDocument]:
        found: list[SourceDocument] = []
        for grade, route, page_url in GO_PAGES:
            soup = self._soup(page_url)
            for anchor in soup.find_all("a", href=True):
                title = clean_text(anchor.get_text(" ", strip=True))
                href = str(anchor["href"])
                lowered = title.casefold()
                if (
                    "/download/GOPRO-" not in href
                    or "leerplan" not in lowered
                    or "wijzigingen" in lowered
                    or "materi" in lowered
                ):
                    continue
                found.append(
                    SourceDocument(
                        provider="GO",
                        title=title,
                        url=urljoin(GO_BASE, href),
                        page_url=page_url,
                        grade=grade,
                        finality=route if "finaliteit" in route else "",
                        stream=route if "stroom" in route else "",
                        discipline=infer_discipline(title),
                        file_type="pdf",
                    )
                )
        logger.info("GO!: %s publieke leerplandocumenten", len(found))
        return found

    def discover_kov(self) -> list[SourceDocument]:
        subjects: list[tuple[str, str, str]] = []
        for grade, tab in KOV_TABS:
            soup = self._soup(f"{KOV_INDEX}?tab={tab}")
            for anchor in soup.find_all("a", href=True):
                title = clean_text(anchor.get_text(" ", strip=True))
                href = str(anchor["href"])
                if not re.match(r"^/(?:I|II|III|VII)(?:-|$)", href, re.I):
                    continue
                subjects.append((grade, title, urljoin(KOV_BASE, href)))

        found: list[SourceDocument] = []
        for grade, subject_title, subject_url in dict.fromkeys(subjects):
            page_url = subject_url.rstrip("/") + "/leerplan"
            soup = self._soup(page_url)
            candidates: list[tuple[str, str]] = []
            for anchor in soup.find_all("a", href=True):
                label = clean_text(anchor.get_text(" ", strip=True))
                href = str(anchor["href"])
                if "/download/content/" not in href or not href.lower().endswith(".docx"):
                    continue
                candidates.append((label, urljoin(KOV_BASE, href)))
            if not candidates:
                continue
            # De leerplanpagina zet de actuele geldige versie bovenaan.
            label, url = candidates[0]
            route_text = f"{subject_title} {label}".casefold()
            finality = next(
                (
                    value
                    for token, value in (
                        ("d/a-finaliteit", "dubbele finaliteit"),
                        ("d-finaliteit", "doorstroomfinaliteit"),
                        ("a-finaliteit", "arbeidsmarktfinaliteit"),
                    )
                    if token in route_text
                ),
                "",
            )
            stream = infer_stream(route_text)
            found.append(
                SourceDocument(
                    provider="KOV",
                    title=subject_title,
                    url=url,
                    page_url=page_url,
                    grade=grade,
                    finality=finality,
                    stream=stream,
                    discipline=infer_discipline(subject_title),
                    file_type="docx",
                )
            )
            if self.limit is not None and len(found) >= self.limit:
                logger.info("KOV: documentlimiet %s bereikt", self.limit)
                return found
        logger.info("KOV: %s actuele publieke Word-leerplannen", len(found))
        return found

    def discover_ovsg(self) -> list[SourceDocument]:
        index = self._soup(OVSG_INDEX)
        pages: list[tuple[str, str]] = []
        for anchor in index.find_all("a", href=True):
            href = str(anchor["href"])
            title = clean_text(anchor.get_text(" ", strip=True))
            if "/leerplannen-secundair-onderwijs/onderwijsdoelen-" in href:
                pages.append((title, urljoin(OVSG_BASE, href)))

        found: list[SourceDocument] = []
        for page_title, page_url in dict.fromkeys(pages):
            soup = self._soup(page_url)
            route = page_title.casefold()
            grade = infer_grade(route)
            finality = next(
                (
                    value
                    for token, value in (
                        ("doorstroom", "doorstroomfinaliteit"),
                        ("dubbele", "dubbele finaliteit"),
                        ("arbeidsmarkt", "arbeidsmarktfinaliteit"),
                    )
                    if token in route
                ),
                "",
            )
            stream = infer_stream(route)
            for anchor in soup.find_all("a", href=True):
                title = clean_text(anchor.get_text(" ", strip=True))
                href = str(anchor["href"])
                lowered = title.casefold()
                if (
                    "actions/protectedlinks/link/get" not in href
                    or "leerplan" not in lowered
                    or "materi" in lowered
                ):
                    continue
                found.append(
                    SourceDocument(
                        provider="OVSG",
                        title=title,
                        url=urljoin(OVSG_BASE, href),
                        page_url=page_url,
                        grade=grade,
                        finality=finality,
                        stream=stream,
                        discipline=infer_discipline(title),
                        file_type="pdf",
                    )
                )
        logger.info("OVSG: %s publieke leerplandocumenten", len(found))
        return found

    def _soup(self, url: str) -> BeautifulSoup:
        response = self.session.get(url, timeout=self.timeout)
        response.raise_for_status()
        return BeautifulSoup(response.text, "html.parser")

    def _load_document(self, source: SourceDocument) -> tuple[bytes, str]:
        suffix = ".docx" if source.file_type == "docx" else ".pdf"
        provider_dir = self.download_dir / source.provider.casefold()
        provider_dir.mkdir(parents=True, exist_ok=True)
        target = provider_dir / safe_name(source.title + source.url, suffix)
        if self.skip_download:
            if not target.exists():
                raise FileNotFoundError(f"Cache ontbreekt: {target}")
            return target.read_bytes(), source.file_type

        # OVSG protected links require the CSRF-cookie from the referring page.
        if source.provider == "OVSG":
            self.session.get(source.page_url, timeout=self.timeout).raise_for_status()
        response = self.session.get(
            source.url,
            headers={"Referer": source.page_url},
            timeout=self.timeout,
        )
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "").casefold()
        file_type = (
            "docx"
            if "wordprocessingml" in content_type or source.url.lower().endswith(".docx")
            else "pdf"
        )
        if file_type == "pdf" and not response.content.startswith(b"%PDF"):
            raise ValueError(f"verwacht PDF, kreeg {content_type or 'onbekend formaat'}")
        target.write_bytes(response.content)
        return response.content, file_type

    def _write(
        self,
        records: list[GoalRecord],
        documents: list[SourceDocument],
        parsed_documents: int,
    ) -> None:
        self.output.parent.mkdir(parents=True, exist_ok=True)
        with self.output.open("w", encoding="utf-8") as handle:
            for record in records:
                handle.write(json.dumps(asdict(record), ensure_ascii=False) + "\n")

        report = {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "source_pages": {
                "GO": [page for _, _, page in GO_PAGES],
                "KOV": [f"{KOV_INDEX}?tab={tab}" for _, tab in KOV_TABS],
                "OVSG": [OVSG_INDEX],
            },
            "providers": list(self.providers),
            "documents_found": len(documents),
            "documents_parsed": parsed_documents,
            "record_count": len(records),
            "provider_counts": dict(Counter(record.netwerk for record in records)),
            "warnings": self.warnings,
        }
        self.report_path.write_text(
            json.dumps(report, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        logger.info("JSONL: %s (%s doelen)", self.output, len(records))

    def _warn(self, message: str) -> None:
        self.warnings.append(message)
        logger.warning(message)

    @staticmethod
    def _dedupe_documents(documents: list[SourceDocument]) -> list[SourceDocument]:
        return list({document.url: document for document in documents}.values())

    @staticmethod
    def _dedupe_records(records: list[GoalRecord]) -> list[GoalRecord]:
        unique: dict[tuple[str, str, str], GoalRecord] = {}
        for record in records:
            unique.setdefault(record.key(), record)
        return list(unique.values())


def extract_goal_sentence(block: str) -> str:
    cleaned = clean_text(block)
    match = LEARNER_SENTENCE_RE.search(cleaned)
    if match:
        return clean_text(match.group(1))
    fallback = re.split(
        r"\b(?:SC\s+\d+|Activeren voorkennis|Ondersteuning|Verdieping|"
        r"GOEDGEKEURD TOT|MD\s+\d)",
        cleaned,
        maxsplit=1,
        flags=re.I,
    )[0].strip(" #·-")
    if 15 <= len(fallback) <= 900:
        return fallback
    return ""


def parse_pdf(payload: bytes, source: SourceDocument) -> list[GoalRecord]:
    reader = PdfReader(BytesIO(payload))
    records: list[GoalRecord] = []
    current_section = ""

    for page in reader.pages:
        text = page.extract_text() or ""
        section_match = re.search(
            r"(?im)^Sleutelcompetentie\s+(\d+)\s*:\s*(.+)$",
            text,
        )
        current_sc_nr = ""
        current_sc_name = current_section
        if section_match:
            current_sc_nr = clean_text(section_match.group(1))
            current_sc_name = clean_text(section_match.group(2))
            current_section = current_sc_name

        matches = list(GOAL_CODE_RE.finditer(text))
        for index, match in enumerate(matches):
            code = clean_text(match.group(1))
            block_end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            block = text[match.end() : block_end]
            title = extract_goal_sentence(block)
            if not title:
                continue
            minimum = MINIMUM_CODE_RE.search(block[:120])
            records.append(
                GoalRecord(
                    code=code,
                    titel=title,
                    discipline=source.discipline or current_sc_name,
                    subdomein=current_section,
                    toelichting="",
                    leerjaar_route=" · ".join(
                        value for value in (source.grade, source.finality, source.stream) if value
                    ),
                    onderwijsniveau="secundair onderwijs",
                    graad=source.grade,
                    finaliteit=source.finality,
                    stroom=source.stream,
                    netwerk=source.provider,
                    bron_url=source.url,
                    bron_titel=source.title,
                    minimumdoel_codes=[minimum.group(1)] if minimum else [],
                    sleutelcompetentie_nr=current_sc_nr,
                    sleutelcompetentie=current_sc_name,
                )
            )
    return records


def _lpds_from_reference(text: str) -> list[int]:
    match = re.search(r"\(LPD\s+([^)]+)\)", text, re.I)
    if not match:
        return []
    values: list[int] = []
    for part in re.split(r"\s*(?:,|en)\s*", match.group(1)):
        range_match = re.match(r"(\d+)\s*(?:t\.?e\.?m\.?|[-–])\s*(\d+)", part, re.I)
        if range_match:
            values.extend(range(int(range_match.group(1)), int(range_match.group(2)) + 1))
        elif part.strip().isdigit():
            values.append(int(part.strip()))
    return values


def parse_kov_docx(payload: bytes, source: SourceDocument) -> list[GoalRecord]:
    document = Document(BytesIO(payload))
    subject_code = urlparse(source.page_url).path.strip("/").split("/")[0]
    minimum_by_lpd: dict[int, list[str]] = {}
    for paragraph in document.paragraphs:
        if paragraph.style.name != "MD + SMD + BK":
            continue
        text = clean_text(paragraph.text)
        code_match = re.match(r"(?:MD|SMD)\s+([A-Z0-9.]+)", text, re.I)
        if not code_match:
            continue
        for number in _lpds_from_reference(text):
            minimum_by_lpd.setdefault(number, []).append(code_match.group(1))

    records: list[GoalRecord] = []
    section = ""
    lpd_number = 0
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if paragraph.style.name == "Heading 2" and text:
            section = text
            continue
        if paragraph.style.name != "Doel" or not text:
            continue
        lpd_number += 1
        records.append(
            GoalRecord(
                code=f"{subject_code} LPD {lpd_number}",
                titel=text,
                discipline=source.discipline,
                subdomein=section,
                toelichting="",
                leerjaar_route=" · ".join(
                    value for value in (source.grade, source.finality, source.stream) if value
                ),
                onderwijsniveau="secundair onderwijs",
                graad=source.grade,
                finaliteit=source.finality,
                stroom=source.stream,
                netwerk=source.provider,
                bron_url=source.url,
                bron_titel=source.title,
                minimumdoel_codes=minimum_by_lpd.get(lpd_number, []),
            )
        )
    return records


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Scrape publieke Vlaamse leerplannen secundair onderwijs."
    )
    parser.add_argument(
        "--provider",
        nargs="+",
        choices=("GO", "KOV", "OVSG"),
        default=["GO", "KOV", "OVSG"],
    )
    parser.add_argument("--output", type=Path, default=OUTPUT_PATH)
    parser.add_argument("--limit", type=int)
    parser.add_argument("--skip-download", action="store_true")
    parser.add_argument("--timeout", type=int, default=90)
    parser.add_argument("-v", "--verbose", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s %(levelname)s %(message)s",
    )
    try:
        SecondaryCurriculumScraper(
            output=args.output,
            providers=args.provider,
            limit=args.limit,
            skip_download=args.skip_download,
            timeout=args.timeout,
        ).run()
        return 0
    except KeyboardInterrupt:
        return 130
    except Exception:
        logger.exception("Scrape mislukt")
        return 1


if __name__ == "__main__":
    sys.exit(main())
